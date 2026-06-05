#!/usr/bin/env python3
"""Create a Korean research-note DOCX from a reusable JSON spec.

The script is intentionally generic:

  python scripts/create_research_note_docx.py \
    --input note.json \
    --output out/research-note.docx \
    [--sample path/to/style-carrier.docx]

`--sample` is optional: if omitted, a blank document is used as the style
carrier. The script clears the body, then writes a research-note document using
the same broad format: opening research-purpose box, numbered prose sections,
list-paragraph bullets, and optional compact tables.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips


FONT = "Malgun Gothic"
BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 31, 31)
GRAY = RGBColor(90, 90, 90)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
BODY_FIRST_LINE_TWIPS = 800
KV_TABLE_WIDTHS_TWIPS = [1701, 7843]
FULL_TABLE_WIDTH_TWIPS = 9544


DATE_PATTERNS = [
    re.compile(r"(20\d{2})\s*[.\-/]\s*(\d{1,2})\s*[.\-/]\s*(\d{1,2})"),
    re.compile(r"(20\d{2})\s*년\s*(\d{1,2})\s*월\s*(\d{1,2})\s*일?"),
]

BRIEF_STYLE_FORBIDDEN = re.compile(
    r"(하였다|했다|되었다|이었다|아니다|것이다|한다|된다|있다|없다|"
    r"보인다|나타났다|발생한다|필요하다|중요하다|의미한다|가능하다|사용한다|"
    r"정리하였다|확인하였다|유지하였다|진행하였다|실행하였다|수행하였다|"
    r"판단하였다|생성하였다|전환하였다|수정하였다)([.。]|$)"
)


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_east_asia_font(run, font_name: str = FONT) -> None:
    run.font.name = font_name
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def parse_date(date_text: str) -> tuple[int, int, int]:
    for pattern in DATE_PATTERNS:
        match = pattern.search(date_text)
        if match:
            return tuple(int(part) for part in match.groups())
    raise ValueError(f"date must include YYYY.MM.DD, YYYY-MM-DD, or YYYY년 M월 D일: {date_text!r}")


def format_date_compact(date_text: str) -> str:
    year, month, day = parse_date(date_text)
    return f"{year:04d}.{month:02d}.{day:02d}"


def format_date_footer(date_text: str) -> str:
    year, month, day = parse_date(date_text)
    return f"{year:04d}. {month:02d}. {day:02d}"


def format_date_korean(date_text: str) -> str:
    year, month, day = parse_date(date_text)
    return f"{year:04d}년 {month}월 {day}일"


def replace_date_placeholders(value: Any, date_text: str) -> Any:
    if isinstance(value, str):
        return (
            value.replace("{date}", format_date_compact(date_text))
            .replace("{footer_date}", format_date_footer(date_text))
            .replace("{date_kr}", format_date_korean(date_text))
        )
    if isinstance(value, list):
        return [replace_date_placeholders(item, date_text) for item in value]
    if isinstance(value, dict):
        return {key: replace_date_placeholders(item, date_text) for key, item in value.items()}
    return value


def iter_spec_strings(value: Any):
    if isinstance(value, str):
        yield value
    elif isinstance(value, list):
        for item in value:
            yield from iter_spec_strings(item)
    elif isinstance(value, dict):
        for item in value.values():
            yield from iter_spec_strings(item)


def assert_brief_note_style(spec: dict[str, Any]) -> None:
    """Reject narrative Korean endings in research-note specs.

    The requested note style is closer to Korean report bullet style: concise
    nominal endings such as `확장함`, `필요함`, `것임`, `아님`.
    """

    violations: list[str] = []
    for text in iter_spec_strings(spec):
        if BRIEF_STYLE_FORBIDDEN.search(text):
            violations.append(text)
    if violations:
        preview = "\n- ".join(violations[:8])
        raise ValueError(f"Research note must use brief report-style endings:\n- {preview}")


def iter_doc_text(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell.text
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                yield paragraph.text
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield cell.text


def assert_no_wrong_dates(doc: Document, date_text: str) -> None:
    expected = parse_date(date_text)
    wrong_dates: list[str] = []
    for text in iter_doc_text(doc):
        for pattern in DATE_PATTERNS:
            for match in pattern.finditer(text):
                found = tuple(int(part) for part in match.groups())
                if found != expected:
                    wrong_dates.append(match.group(0))
    if wrong_dates:
        unique = ", ".join(sorted(set(wrong_dates)))
        raise ValueError(f"Document contains date(s) other than the requested note date: {unique}")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "C8CDD2")


def set_cell_width(cell, width_twips: int) -> None:
    cell.width = Twips(width_twips)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_twips))


def set_table_widths(table, widths_twips: list[int]) -> None:
    table.autofit = False
    table.allow_autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_twips):
                set_cell_width(cell, widths_twips[idx])


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.5, color=DARK) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(str(text))
    run.bold = bold
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    paragraph.add_run()._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)

    placeholder = paragraph.add_run("1")
    placeholder.bold = True
    set_east_asia_font(placeholder)
    placeholder.font.size = Pt(9.5)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def add_footer_run(paragraph, text: str, *, bold: bool = True, size: float = 9.5, color=DARK):
    run = paragraph.add_run(text)
    run.bold = bold
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    size: float = 10.5,
    color=DARK,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.15,
    first_line_twips: int | None = None,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_twips is not None:
        p.paragraph_format.first_line_indent = Twips(first_line_twips)
    run = p.add_run(str(text))
    run.bold = bold
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_title(doc: Document, spec: dict[str, Any]) -> None:
    title = add_paragraph(doc, spec.get("title", "연구노트"), bold=True, size=20, color=BLUE, after=2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_text = spec.get("subtitle")
    if subtitle_text:
        subtitle = add_paragraph(doc, subtitle_text, size=11.5, color=GRAY, after=10)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, KV_TABLE_WIDTHS_TWIPS)
    set_table_borders(table)
    for idx, (key, value) in enumerate(rows):
        set_cell_shading(table.cell(idx, 0), LIGHT_GRAY)
        set_cell_text(table.cell(idx, 0), key, bold=True, size=9.5, color=BLUE)
        set_cell_text(table.cell(idx, 1), value, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_summary_box(doc: Document, purpose: str, summary: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [FULL_TABLE_WIDTH_TWIPS])
    set_table_borders(table)
    set_cell_shading(table.cell(0, 0), LIGHT_BLUE)
    set_cell_text(table.cell(0, 0), f"[연구목적] {purpose}", bold=True, size=9.7, color=BLUE)
    set_cell_text(table.cell(1, 0), f"[연구내용] {summary}", size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_bullet(doc: Document, item: Any) -> None:
    if isinstance(item, dict):
        lead = str(item.get("lead", ""))
        body = str(item.get("body", ""))
    else:
        lead = ""
        body = str(item)

    p = doc.add_paragraph()
    p.style = doc.styles["List Paragraph"]
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12

    if lead:
        run = p.add_run(lead)
        run.bold = True
        set_east_asia_font(run)
        run.font.size = Pt(10.2)
        run.font.color.rgb = DARK
        run = p.add_run(" " + body)
    else:
        run = p.add_run(body)
    set_east_asia_font(run)
    run.font.size = Pt(10.2)
    run.font.color.rgb = DARK


def add_simple_table(doc: Document, table_spec: dict[str, Any]) -> None:
    headers = table_spec.get("headers", [])
    rows = table_spec.get("rows", [])
    if not headers or not rows:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = table_spec.get("widths_twips") or []
    if not widths:
        widths_cm = table_spec.get("widths_cm") or []
        widths = [round(float(width) * 567) for width in widths_cm]
    if not widths:
        widths = [FULL_TABLE_WIDTH_TWIPS // len(headers)] * len(headers)
    set_table_widths(table, [int(width) for width in widths[: len(headers)]])
    set_table_borders(table)

    for c_idx, header in enumerate(headers):
        set_cell_shading(table.cell(0, c_idx), LIGHT_BLUE)
        set_cell_text(table.cell(0, c_idx), header, bold=True, size=9.0, color=BLUE)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row[: len(headers)]):
            set_cell_text(table.cell(r_idx, c_idx), str(text), size=8.7)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_section(doc: Document, section: dict[str, Any]) -> None:
    heading = str(section.get("heading", "")).strip()
    if heading:
        # Match the sample: headings remain Normal paragraphs, bold, numbered.
        size = 14 if "." in heading[:3] and "-" not in heading[:5] else 12.5
        add_paragraph(doc, heading, bold=True, size=size, color=DARK, before=5, after=5)

    for paragraph in section.get("paragraphs", []):
        if isinstance(paragraph, dict):
            text = str(paragraph.get("text", ""))
            if paragraph.get("dash") and not text.lstrip().startswith("-"):
                text = "- " + text
            first_line = paragraph.get("first_line_twips", BODY_FIRST_LINE_TWIPS)
        else:
            text = str(paragraph)
            first_line = BODY_FIRST_LINE_TWIPS
        add_paragraph(doc, text, size=10.5, after=6, first_line_twips=int(first_line))

    for bullet in section.get("bullets", []):
        add_bullet(doc, bullet)

    for table_spec in section.get("tables", []):
        add_simple_table(doc, table_spec)


def add_footer(doc: Document, spec: dict[str, Any]) -> None:
    date = format_date_footer(str(spec.get("date", "")).strip())
    researcher = str(spec.get("footer_author") or spec.get("researcher", "")).split("/")[0].strip()
    reviewer = str(spec.get("footer_reviewer") or researcher).strip()
    company = str(spec.get("footer_company") or "").strip()

    for section in doc.sections:
        footer = section.footer
        for child in list(footer._element):
            footer._element.remove(child)

        signature = footer.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_footer_run(signature, "페이지 ", bold=False)
        add_field(signature, "PAGE")
        add_footer_run(signature, " / ", bold=False)
        add_field(signature, "NUMPAGES")
        add_footer_run(signature, f" {date} {company} 작성자 : {researcher} _____________ / 검토자 : {reviewer} _____________")


def build_docx(spec: dict[str, Any], output: Path, sample: Path | None = None) -> None:
    spec = replace_date_placeholders(spec, str(spec.get("date", "")))
    assert_brief_note_style(spec)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(sample)) if sample else Document()
    clear_body(doc)

    section = doc.sections[0]
    section.top_margin = Inches(float(spec.get("top_margin_in", 0.82)))
    section.bottom_margin = Inches(float(spec.get("bottom_margin_in", 0.82)))
    section.left_margin = Inches(float(spec.get("left_margin_in", 0.82)))
    section.right_margin = Inches(float(spec.get("right_margin_in", 0.82)))
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)

    add_title(doc, spec)
    add_kv_table(
        doc,
        [
            ("과제명", spec.get("project_name", "")),
            ("연구자", spec.get("researcher", "")),
            ("작성일", spec.get("date", "")),
            ("연구 구분", spec.get("research_type", "")),
            ("연구 저장소", spec.get("research_storage", "")),
        ],
    )
    add_summary_box(doc, spec.get("purpose", ""), spec.get("summary", ""))

    for section_spec in spec.get("sections", []):
        add_section(doc, section_spec)

    for table_spec in spec.get("tables", []):
        add_simple_table(doc, table_spec)

    add_footer(doc, spec)
    assert_no_wrong_dates(doc, str(spec.get("date", "")))

    props = doc.core_properties
    props.title = spec.get("core_title", spec.get("title", "연구노트"))
    props.subject = spec.get("subtitle", "")
    props.author = spec.get("researcher", "")
    props.comments = "Generated by agent-rnd-phd/scripts/create_research_note_docx.py"

    doc.save(str(output))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="Research note JSON spec")
    parser.add_argument(
        "--sample", type=Path, default=None,
        help="Optional DOCX style carrier. If omitted, a blank document is used.",
    )
    parser.add_argument("--output", required=True, type=Path, help="Output DOCX path")
    args = parser.parse_args()

    spec = json.loads(args.input.read_text(encoding="utf-8"))
    build_docx(spec, args.output, args.sample)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
