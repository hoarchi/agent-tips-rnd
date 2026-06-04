#!/usr/bin/env python3
"""Create an academic manuscript DOCX from a reusable JSON spec.

Sibling generator to ``create_research_note_docx.py``. Where the research-note
generator emits a Korean R&D 연구노트 (brief report style), this generator emits
a standard academic paper (IMRaD): title, authors, affiliations, abstract,
keywords, multi-level numbered sections with narrative prose, numbered figures
and tables with captions, optional equations, and a numbered reference list.

  python scripts/create_manuscript_docx.py \
    --input templates/manuscript.json \
    --output out/manuscript.docx \
    [--sample path/to/style-carrier.docx]

Unlike the research note, this generator does NOT enforce brief nominal endings.
Papers are written in full narrative prose, in Korean or English, in whichever
language the spec headings and body are written. Section levels, figure/table
numbering, and references are handled automatically.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips


# Academic default: a serif latin face with a Korean east-asian fallback reads
# closest to a journal/conference template while still rendering Hangul cleanly.
FONT_EA = "Malgun Gothic"
FONT_LATIN = "Times New Roman"
DARK = RGBColor(26, 26, 26)
GRAY = RGBColor(90, 90, 90)
ACCENT = RGBColor(31, 78, 121)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "EAF1F8"
FULL_WIDTH_TWIPS = 9544
BODY_FIRST_LINE_TWIPS = 480

# Leading numeric token of a heading, e.g. "2", "2.1", "3.2.4".
HEADING_NUMBER = re.compile(r"^\s*(\d+(?:\.\d+)*)")


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_fonts(run, *, latin: str | None = None, east_asia: str | None = None) -> None:
    latin = latin or FONT_LATIN
    east_asia = east_asia or FONT_EA
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    italic: bool = False,
    size: float = 10.5,
    color=DARK,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.3,
    align=None,
    first_line_twips: int | None = None,
    fonts: tuple[str, str] | None = None,
):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    if first_line_twips is not None:
        pf.first_line_indent = Twips(first_line_twips)
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    latin, east_asia = fonts or (FONT_LATIN, FONT_EA)
    set_fonts(run, latin=latin, east_asia=east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = "C8CDD2") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_twips: int) -> None:
    cell.width = Twips(width_twips)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_twips))


def set_table_widths(table, widths_twips: list[int]) -> None:
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_twips):
                set_cell_width(cell, widths_twips[idx])


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: float = 9.0,
    color=DARK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    set_fonts(run)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def heading_level(heading: str, explicit: int | None) -> int:
    if explicit:
        return max(1, min(3, int(explicit)))
    match = HEADING_NUMBER.match(heading)
    if match:
        depth = match.group(1).count(".") + 1
        return max(1, min(3, depth))
    return 1


def add_heading(doc: Document, heading: str, level: int) -> None:
    size = {1: 12.5, 2: 11.0, 3: 10.5}[level]
    italic = level == 3
    add_paragraph(
        doc,
        heading,
        bold=True,
        italic=italic,
        size=size,
        color=DARK,
        before={1: 11, 2: 8, 3: 6}[level],
        after={1: 5, 2: 4, 3: 3}[level],
        line_spacing=1.2,
    )


def add_title_block(doc: Document, spec: dict[str, Any]) -> None:
    add_paragraph(
        doc,
        spec.get("title", "Untitled Manuscript"),
        bold=True,
        size=18,
        color=DARK,
        after=3,
        line_spacing=1.18,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    subtitle = spec.get("subtitle")
    if subtitle:
        add_paragraph(
            doc, subtitle, size=12.5, color=GRAY, after=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    authors = spec.get("authors", [])
    if authors:
        parts = []
        for author in authors:
            if isinstance(author, dict):
                name = str(author.get("name", "")).strip()
                mark = str(author.get("mark", "")).strip()
                parts.append(name + mark)
            else:
                parts.append(str(author).strip())
        add_paragraph(
            doc, ", ".join(p for p in parts if p), size=11, color=DARK,
            after=2, align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for affiliation in spec.get("affiliations", []):
        add_paragraph(
            doc, str(affiliation), size=9.5, color=GRAY, after=1,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    corresponding = spec.get("corresponding")
    if corresponding:
        add_paragraph(
            doc, str(corresponding), size=9.5, color=GRAY, after=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_abstract(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [FULL_WIDTH_TWIPS])
    set_table_borders(table)
    set_cell_shading(table.cell(0, 0), LIGHT_BLUE)

    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(f"{label}  ")
    lead.bold = True
    set_fonts(lead)
    lead.font.size = Pt(9.7)
    lead.font.color.rgb = ACCENT
    run = p.add_run(str(text))
    set_fonts(run)
    run.font.size = Pt(9.7)
    run.font.color.rgb = DARK
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_keywords(doc: Document, label: str, keywords: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    lead = p.add_run(f"{label}  ")
    lead.bold = True
    set_fonts(lead)
    lead.font.size = Pt(9.5)
    lead.font.color.rgb = ACCENT
    run = p.add_run(", ".join(str(k) for k in keywords))
    run.italic = True
    set_fonts(run)
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK


def add_body_paragraph(doc: Document, paragraph: Any) -> None:
    if isinstance(paragraph, dict):
        text = str(paragraph.get("text", ""))
        indent = paragraph.get("first_line_twips", BODY_FIRST_LINE_TWIPS)
        no_indent = paragraph.get("no_indent", False)
    else:
        text = str(paragraph)
        indent = BODY_FIRST_LINE_TWIPS
        no_indent = False
    add_paragraph(
        doc, text, size=10.5, after=6, line_spacing=1.3,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_twips=None if no_indent else int(indent),
    )


def add_bullet(doc: Document, item: Any) -> None:
    if isinstance(item, dict):
        lead = str(item.get("lead", ""))
        body = str(item.get("body", ""))
    else:
        lead = ""
        body = str(item)
    p = doc.add_paragraph()
    p.style = doc.styles["List Paragraph"]
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if lead:
        run = p.add_run(lead + " ")
        run.bold = True
        set_fonts(run)
        run.font.size = Pt(10.2)
        run.font.color.rgb = DARK
        run = p.add_run(body)
    else:
        run = p.add_run(body)
    set_fonts(run)
    run.font.size = Pt(10.2)
    run.font.color.rgb = DARK


def add_table_block(doc: Document, table_spec: dict[str, Any], counters: dict[str, int], label_prefix: str) -> None:
    headers = table_spec.get("headers", [])
    rows = table_spec.get("rows", [])
    if not headers or not rows:
        return

    caption = table_spec.get("caption")
    if caption:
        counters["table"] += 1
        add_paragraph(
            doc, f"{label_prefix} {counters['table']}. {caption}",
            bold=True, size=9.2, color=DARK, before=4, after=3,
            align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15,
        )

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = table_spec.get("widths_twips") or []
    if not widths:
        widths_cm = table_spec.get("widths_cm") or []
        widths = [round(float(w) * 567) for w in widths_cm]
    if not widths:
        widths = [FULL_WIDTH_TWIPS // len(headers)] * len(headers)
    set_table_widths(table, [int(w) for w in widths[: len(headers)]])
    set_table_borders(table)

    for c_idx, header in enumerate(headers):
        set_cell_shading(table.cell(0, c_idx), LIGHT_GRAY)
        set_cell_text(
            table.cell(0, c_idx), header, bold=True, size=9.0, color=ACCENT,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row[: len(headers)]):
            set_cell_text(table.cell(r_idx, c_idx), str(text), size=8.7)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_figure(doc: Document, figure_spec: dict[str, Any], counters: dict[str, int], label_prefix: str, base_dir: Path) -> None:
    counters["figure"] += 1
    image = figure_spec.get("image")
    width_in = float(figure_spec.get("width_in", 5.5))
    placed = False
    if image:
        image_path = (base_dir / image) if not Path(image).is_absolute() else Path(image)
        if image_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.add_run().add_picture(str(image_path), width=Inches(width_in))
            placed = True
    if not placed:
        # Reserved placeholder box so the layout slot is visible before the
        # final figure asset exists.
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_widths(table, [int(width_in * 1440)])
        set_table_borders(table, color="B0B6BC")
        set_cell_shading(table.cell(0, 0), LIGHT_GRAY)
        placeholder = figure_spec.get("placeholder") or f"[figure asset pending: {image or 'no image path'}]"
        set_cell_text(table.cell(0, 0), placeholder, italic=True, size=9.0, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    caption = figure_spec.get("caption", "")
    add_paragraph(
        doc, f"{label_prefix} {counters['figure']}. {caption}",
        bold=True, size=9.2, color=DARK, before=2, after=6,
        align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15,
    )


def add_equation(doc: Document, eq_spec: Any, counters: dict[str, int]) -> None:
    if isinstance(eq_spec, dict):
        text = str(eq_spec.get("text", ""))
        numbered = bool(eq_spec.get("number", True))
    else:
        text = str(eq_spec)
        numbered = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    set_fonts(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    if numbered:
        counters["equation"] += 1
        tab = p.add_run(f"\t({counters['equation']})")
        set_fonts(tab)
        tab.font.size = Pt(10.5)
        tab.font.color.rgb = DARK


def add_section(doc: Document, section: dict[str, Any], counters: dict[str, int], spec: dict[str, Any], base_dir: Path) -> None:
    heading = str(section.get("heading", "")).strip()
    if heading:
        add_heading(doc, heading, heading_level(heading, section.get("level")))

    for paragraph in section.get("paragraphs", []):
        add_body_paragraph(doc, paragraph)

    for bullet in section.get("bullets", []):
        add_bullet(doc, bullet)

    for equation in section.get("equations", []):
        add_equation(doc, equation, counters)

    for figure in section.get("figures", []):
        add_figure(doc, figure, counters, spec.get("figure_label_prefix", "Figure"), base_dir)

    for table_spec in section.get("tables", []):
        add_table_block(doc, table_spec, counters, spec.get("table_label_prefix", "Table"))

    for child in section.get("sections", []):
        add_section(doc, child, counters, spec, base_dir)


def add_references(doc: Document, label: str, references: list[Any]) -> None:
    if not references:
        return
    add_heading(doc, label, 1)
    for idx, reference in enumerate(references, start=1):
        text = reference if isinstance(reference, str) else str(reference.get("text", ""))
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(3)
        pf.line_spacing = 1.2
        pf.left_indent = Twips(420)
        pf.first_line_indent = Twips(-420)
        run = p.add_run(f"[{idx}] {text}")
        set_fonts(run)
        run.font.size = Pt(9.3)
        run.font.color.rgb = DARK


def build_docx(spec: dict[str, Any], output: Path, sample: Path | None, base_dir: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(sample)) if sample else Document()
    clear_body(doc)

    section = doc.sections[0]
    section.top_margin = Inches(float(spec.get("top_margin_in", 0.9)))
    section.bottom_margin = Inches(float(spec.get("bottom_margin_in", 0.9)))
    section.left_margin = Inches(float(spec.get("left_margin_in", 0.95)))
    section.right_margin = Inches(float(spec.get("right_margin_in", 0.95)))

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_EA)
    normal.font.size = Pt(10.5)

    counters = {"figure": 0, "table": 0, "equation": 0}

    add_title_block(doc, spec)

    if spec.get("abstract"):
        add_abstract(doc, spec.get("abstract_label", "Abstract"), spec["abstract"])
    if spec.get("abstract_secondary"):
        add_abstract(doc, spec.get("abstract_secondary_label", "초록"), spec["abstract_secondary"])
    if spec.get("keywords"):
        add_keywords(doc, spec.get("keywords_label", "Keywords —"), spec["keywords"])

    for section_spec in spec.get("sections", []):
        add_section(doc, section_spec, counters, spec, base_dir)

    for table_spec in spec.get("tables", []):
        add_table_block(doc, table_spec, counters, spec.get("table_label_prefix", "Table"))

    add_references(doc, spec.get("references_label", "References"), spec.get("references", []))

    props = doc.core_properties
    props.title = spec.get("title", "Manuscript")
    props.subject = spec.get("subtitle", "")
    props.author = "; ".join(
        (a.get("name", "") if isinstance(a, dict) else str(a)) for a in spec.get("authors", [])
    )
    props.comments = "Generated by agent-rnd-phd/scripts/create_manuscript_docx.py"

    doc.save(str(output))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="Manuscript JSON spec")
    parser.add_argument("--output", required=True, type=Path, help="Output DOCX path")
    parser.add_argument(
        "--sample", type=Path, default=None,
        help="Optional DOCX style carrier. If omitted, a blank document is used.",
    )
    args = parser.parse_args()

    spec = json.loads(args.input.read_text(encoding="utf-8"))
    build_docx(spec, args.output, args.sample, args.input.resolve().parent)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
